from os import close
import docx
       
#Путь к документу
filepath = input( 'Укажите полный путь к файлу включая имя файла и расширение: ' )
# Присваиваем переменной r путь к файлу
r = ( filepath )
# Вводим переменную с путем
doc = docx.Document(r)
all_paras = doc.paragraphs
len(all_paras)
# Присваиваем переменной а искомое слово
a = input("Введите искомое заменяемое слово:")
b = input("Введите заменяющее слово:")
for p in doc.paragraphs:
    # a - это наше искомое слово которое нужно заменить
        if a in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if a in inline[i].text:
                    # 'new text' это слово на которое будет заменено наше искомое слово
                    text = inline[i].text.replace(a, b)
                    inline[i].text = text
                    print
                    p.text
# Сохранение файла
doc.save(r)
d = input("Замена выполнена, желаете вывести результат в консоль?\n 0 - Нет\n 1 - Да\n")
if d == "1":    
    # Вывод результата в консоль                    
    for para in all_paras:
        print(para.text)
else:
    close

